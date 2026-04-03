"""
=========================================================================================
PROJE: Kurumsal Doküman (Fatura & Teklif) Oluşturucu Platform
AÇIKLAMA: Streamlit tabanlı, dinamik tablo yönetimi yapabilen, otomatik matematiksel
hesaplamalar içeren ve anlık olarak PDF, Word, Excel formatlarında kurumsal antetli
çıktılar üreten web uygulamasıdır. İçerisinde durum yönetimi (session_state) ve 
JSON tabanlı yedekleme/geri yükleme (Save/Load) mimarisi barındırır.
=========================================================================================
"""

import streamlit as st
import pandas as pd
import io
import datetime
import os
import json

# PDF Üretimi için FPDF kütüphanesi
from fpdf import FPDF

# Word Üretimi için python-docx modülleri
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Excel Üretimi için openpyxl modülleri
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as xlImage

# ---------------------------------------------------------------------------------------
# APP KONFİGÜRASYONU
# Sayfa yapısını geniş formatta ve varsayılan menü açık şekilde başlatır.
# ---------------------------------------------------------------------------------------
st.set_page_config(
    layout="wide",
    page_title="Doküman Oluşturucu Platform",
    initial_sidebar_state="expanded"
)

# =========================================================
# SİSTEM SABİTLERİ
# =========================================================
ANTET_DOSYASI = "antet.png"              # PDF ve Excel çıktıları için arka plan/logo dosyası
WORD_TEMPLATE = "word_template.docx"     # Word çıktısı için temel şablon dosyası

# =========================================================
# YARDIMCI FONKSİYONLAR (HELPER METHODS)
# =========================================================

def cevir_tr(metin):
    """
    FPDF kütüphanesi varsayılan olarak Türkçe karakter (UTF-8) desteklemediği için,
    PDF dökümanında hata almamak adına Türkçe karakterleri İngilizce karşılıklarına çevirir.
    """
    tr_map = {
        'ş': 's', 'Ş': 'S',
        'ı': 'i', 'İ': 'I',
        'ğ': 'g', 'Ğ': 'G',
        'ü': 'u', 'Ü': 'U',
        'ö': 'o', 'Ö': 'O',
        'ç': 'c', 'Ç': 'C'
    }
    metin = str(metin)
    for k, v in tr_map.items():
        metin = metin.replace(k, v)
    return metin

def get_alignment(col_name):
    """
    Sütun başlığına bakarak hücre içindeki metnin hizalamasını dinamik olarak belirler.
    Finansal veriler sağa (R), miktarlar/sıra numaraları ortaya (C), metinler sola (L) hizalanır.
    """
    name = str(col_name).lower()
    if any(x in name for x in ['fiyat', 'price', 'tutar', 'total', 'amount']):
        return 'R'
    if any(x in name for x in ['sıra', 'no', 'kdv', 'adet', 'unit', 'qty', 'miktar']):
        return 'C'
    return 'L'

def get_pdf_widths(headers, total_w=190):
    """
    PDF tablosunun A4 kağıda tam sığması için (toplam genişlik genelde 190mm), 
    sütun isimlerine göre oransal genişlik hesaplar.
    """
    widths = []
    for h in headers:
        name = str(h).lower()
        if any(x in name for x in ['sıra', 'no']):
            widths.append(10)
        elif any(x in name for x in ['kdv', 'adet', 'unit', 'qty']):
            widths.append(15)
        elif any(x in name for x in ['fiyat', 'price', 'tutar', 'amount', 'total']):
            widths.append(25)
        elif any(x in name for x in ['marka', 'brand']):
            widths.append(25)
        elif any(x in name for x in ['açıklama', 'remark', 'işlem', 'description']):
            widths.append(60)
        else:
            widths.append(25)

    # Toplam genişliği orantılayarak tabloyu kağıda yayar
    total = sum(widths)
    scale = total_w / total if total > 0 else 1
    return [w * scale for w in widths]

def set_excel_col_widths(ws, headers):
    """
    Excel tablosundaki hücrelerin okunabilir olması için sütun içeriklerine göre
    genişlik değerlerini (width) statik olarak ayarlar.
    """
    for i, header in enumerate(headers, 1):
        col_letter = get_column_letter(i)
        name = str(header).lower()

        if any(x in name for x in ['sıra', 'no', 'kdv', 'adet', 'unit', 'qty']):
            ws.column_dimensions[col_letter].width = 8
        elif any(x in name for x in ['fiyat', 'price', 'tutar', 'amount', 'total']):
            ws.column_dimensions[col_letter].width = 16
        elif any(x in name for x in ['marka', 'brand']):
            ws.column_dimensions[col_letter].width = 15
        elif any(x in name for x in ['açıklama', 'remark', 'işlem', 'description']):
            ws.column_dimensions[col_letter].width = 40
        else:
            ws.column_dimensions[col_letter].width = 20

def get_birim_col(df_columns):
    """
    Tablo içinde 'Birim Fiyat' sütununu bulur. Fiyat gizleme işlemi için gereklidir.
    """
    for col in df_columns:
        if "birim fiyat" in str(col).lower():
            return col
    return None

def format_money_value(val, kur_m):
    """
    Finansal değerleri binlik ayraçlı ve para birimli formata dönüştürür.
    Hatalı veya sıfır değerler için '-NIL-' döndürür. Örn: 1000 -> 1.000 TL
    """
    try:
        fiyat = float(str(val).replace(',', '.'))
        if pd.isna(fiyat) or fiyat <= 0:
            return "-NIL-"
        return f"{fiyat:,.0f}".replace(",", ".") + f" {kur_m}"
    except Exception:
        return "-NIL-"

def guvenli_table_style(table, style_name="Table Grid"):
    """
    Word şablonunda belirli bir tablo stili yoksa uygulamanın çökmesini engellemek için
    yazılmış güvenli stil atama fonksiyonudur (Try-Catch bloğu içerir).
    """
    try:
        table.style = style_name
    except Exception:
        pass

def kolonu_bul(columns, adaylar):
    """
    Kullanıcının sütun adını büyük/küçük veya hatalı yazma ihtimaline karşı
    belirlenen anahtar kelimelere (adaylar) göre doğru sütunu tespit eder.
    """
    for col in columns:
        low = str(col).strip().lower()
        if low in adaylar:
            return col
    return None

def otomatik_hesaplari_uygula(dataframe, sablon_tipi):
    """
    MERKEZİ HESAPLAMA MOTORU:
    1. Yeni veya boş satırlardaki KDV oranlarını otomatik olarak %20'ye sabitler.
    2. Adet, Miktar vb. sütunlar ile Birim Fiyat sütunlarını algılayıp çarparak Tutar'ı otomatik hesaplar.
    3. Null veya NaN (boş) değerleri matematiksel işleme uygun hale getirir.
    """
    df_calc = dataframe.copy()

    # NaN ve format düzeltmeleri
    for col in df_calc.columns:
        low = str(col).strip().lower()
        if low in {"adet", "qty", "quantity", "miktar"}:
            df_calc[col] = pd.to_numeric(df_calc[col], errors="coerce").fillna(0)
        elif low in {"birim fiyatı", "birim fiyati", "birim fiyat", "unit price", "tutar", "total", "amount"}:
            df_calc[col] = pd.to_numeric(df_calc[col], errors="coerce").fillna(0.0)
        elif low == "kdv" and sablon_tipi == "📄 Standart Proforma Fatura":
            df_calc[col] = df_calc[col].apply(lambda x: "%20" if pd.isna(x) or str(x).strip() == "" or str(x).strip().lower() == "nan" else x)
        else:
            df_calc[col] = df_calc[col].fillna("")

    # Çarpım İşlemi (Adet * Birim Fiyat = Tutar)
    if sablon_tipi == "📄 Standart Proforma Fatura":
        adet_col = kolonu_bul(df_calc.columns, {"adet", "qty", "quantity", "miktar"})
        birim_col = kolonu_bul(df_calc.columns, {"birim fiyatı", "birim fiyati", "birim fiyat", "unit price"})
        tutar_col = kolonu_bul(df_calc.columns, {"tutar", "total", "amount"})

        if adet_col is not None and birim_col is not None and tutar_col is not None:
            adet_seri = pd.to_numeric(df_calc[adet_col], errors="coerce").fillna(0)
            birim_seri = pd.to_numeric(df_calc[birim_col], errors="coerce").fillna(0.0)
            df_calc[tutar_col] = (adet_seri * birim_seri).astype(float)

    return df_calc

def toplam_sutununu_bul(dataframe, sablon_tipi):
    """
    Genel hesaplamalar için tutarların bulunduğu (toplanacak) ana sütunu tespit eder.
    """
    if sablon_tipi == "📄 Standart Proforma Fatura":
        tutar_col = kolonu_bul(dataframe.columns, {"tutar", "total", "amount"})
        if tutar_col is not None:
            return tutar_col
    return dataframe.columns[-1]

# =========================================================
# YÖNETİM PANELİ VE YEDEKLEME SİSTEMİ (SIDEBAR)
# =========================================================
st.sidebar.markdown("### ⚙️ Sistem Ayarları")

# Şablon Seçimi
secili_sablon = st.sidebar.radio(
    "📝 Çalışma Şablonunu Seçin:",
    ["⚓ INNOMAR Özel Teklif", "📄 Standart Proforma Fatura"]
)

# Gizlilik Ayarı
gizle_checkbox = st.sidebar.checkbox(
    "🔒 Birim Fiyat Sütununu Çıktılarda Gizle",
    value=False,
    help="İşaretlendiğinde, indirilen dosyalarda Birim Fiyat sütunu tamamen kaldırılır. (Sitede görünmeye devam eder)."
)

# --- TASLAK GERİ YÜKLEME MİMARİSİ ---
st.sidebar.markdown("---")
st.sidebar.markdown("### 💾 Yedekleme & Geri Yükleme")
st.sidebar.caption("Taslaklarınızı bilgisayarınıza kaydedip daha sonra geri yükleyebilirsiniz.")

yuklenen_dosya = st.sidebar.file_uploader("📂 Taslak Yükle (.json)", type=["json"])
if yuklenen_dosya is not None:
    if st.sidebar.button("🔄 Yüklü Taslağı Sisteme Uygula", type="primary"):
        try:
            yedek_veri = json.load(yuklenen_dosya)
            
            # Yüklenen JSON verilerini Session State (Oturum Belleği) içine aktarıyoruz
            st.session_state.aktif_sablon = yedek_veri.get("sablon", "⚓ INNOMAR Özel Teklif")
            st.session_state.not_alani = yedek_veri.get("notlar", "")
            st.session_state.teklif_basligi_str = yedek_veri.get("teklif_basligi", "MY ADA DRY DOCK SERVICES QUOTATION;")
            
            # Veri tablosunu (Dataframe) yeniden inşa ediyoruz
            df_icerik = pd.DataFrame(yedek_veri.get("df", []))
            if not df_icerik.empty:
                st.session_state.veri_df = df_icerik
            
            st.sidebar.success("Taslak başarıyla yüklendi!")
            st.rerun() # Arayüzü yeni verilerle yenile
        except Exception as e:
            st.sidebar.error("Hatalı dosya formatı!")

# İlk açılışta veya şablon değiştiğinde varsayılan verileri sisteme yükler
if 'aktif_sablon' not in st.session_state or st.session_state.aktif_sablon != secili_sablon:
    st.session_state.aktif_sablon = secili_sablon

    if secili_sablon == "⚓ INNOMAR Özel Teklif":
        data = {
            'INSPECTION REMARK': ['ANA MAKİNE BAKIMLARI', 'SU YAPICI BAKIMLARI'],
            'UNIT': ['2 PIECES', '1 SET'],
            'PRICE': [40000.0, 12000.0]
        }
        st.session_state.not_alani = """* IMPORTANT NOTICE;
- DURING MAINTENANCE IF DEFORMATION DETECTED ON WORKING SURFACE AND NEEDED TO RENEW COMPONENTS EACH PARTS WILL BE PRICED ADDITIONALLY.

* REMARKS;
- DELIVERY TIME FOR THE JOB IS 35 DAYS,
- A DETAILED REPORT WILL BE SUBMITTED TO YOUR SIDE UPON COMPLETION OF THE WORK,
- PAYMENT WILL BE ACCEPTED AS BELOW;
    - %50 BEFORE WORK BEGINS,
    - %50 UPON COMPLETION OF THE WORK."""
    else:
        data = {
            'Açıklama': ['Örnek Hizmet', ''],
            'Adet': [1, 2],
            'KDV': ['%20', '%20'],
            'Birim Fiyatı': [1000.0, 500.0],
            'Tutar': [1000.0, 1000.0]
        }
        st.session_state.not_alani = ""

    st.session_state.veri_df = pd.DataFrame(data)
    st.rerun()

st.markdown(
    f"<h2 style='text-align: center;'>{secili_sablon.upper()} SİSTEMİ</h2>",
    unsafe_allow_html=True
)

# =========================================================
# ÜST PANEL: TARİH VE PARA BİRİMİ YÖNETİMİ
# =========================================================
col_t, col_kur = st.columns([1, 1])

secilen_tarih = col_t.date_input("Belge Tarihi", datetime.date.today())
tarih_metni = secilen_tarih.strftime("%d.%m.%Y")
dosya_tarihi = secilen_tarih.strftime("%d_%m_%Y")

kur_secimi = col_kur.selectbox("Para Birimi", ["Euro (€)", "Dolar ($)", "Türk Lirası (₺)"])

# Kur seçimine göre belge sembollerini ve metinlerini ayarlar
if "Euro" in kur_secimi:
    sembol, kur_metin = "€", "EURO"
elif "Dolar" in kur_secimi:
    sembol, kur_metin = "$", "USD"
else:
    sembol, kur_metin = "₺", "TL"

# Özel Teklif şablonu için Dinamik Başlık Alanı
teklif_basligi_str = ""
if secili_sablon == "⚓ INNOMAR Özel Teklif":
    st.write("---")
    varsayilan_baslik = st.session_state.get("teklif_basligi_str", "MY ADA DRY DOCK SERVICES QUOTATION;")
    teklif_basligi_str = st.text_input("⚓ Teklif Başlığı:", varsayilan_baslik)
    st.session_state.teklif_basligi_str = teklif_basligi_str

# =========================================================
# DİNAMİK SÜTUN YÖNETİMİ (KOLON EKLEME/ÇIKARMA)
# =========================================================
st.write("---")
st.caption(
    "Aşağıdaki kutuya virgülle ayırarak istediğiniz kadar sütun ekleyebilir veya silebilirsiniz. "
    "**DİKKAT: Hesaplamaların doğru çalışması için fiyat/tutar sütunu her zaman EN SONDA olmalıdır.**"
)

mevcut_sutunlar = ", ".join(st.session_state.veri_df.columns)
yeni_sutunlar_str = st.text_input("Tablo Sütunları:", mevcut_sutunlar)

yeni_sutunlar = [s.strip() for s in yeni_sutunlar_str.split(",") if s.strip()]
yeni_sutunlar = list(dict.fromkeys(yeni_sutunlar))

if len(yeni_sutunlar) < 2:
    st.warning("Lütfen tabloda en az 2 sütun bırakın.")
    yeni_sutunlar = st.session_state.veri_df.columns.tolist()

# Eğer kullanıcı sütun isimlerini değiştirirse, eski verileri kaybetmeden yeni tabloyu inşa eder
if yeni_sutunlar != st.session_state.veri_df.columns.tolist():
    eski_df = st.session_state.veri_df
    yeni_df = pd.DataFrame(columns=yeni_sutunlar)

    for col in yeni_sutunlar:
        if col in eski_df.columns:
            yeni_df[col] = eski_df[col]
        else:
            yeni_df[col] = ""

    son_sutun_adi = yeni_sutunlar[-1]
    yeni_df[son_sutun_adi] = pd.to_numeric(yeni_df[son_sutun_adi], errors='coerce').fillna(0.0)

    st.session_state.veri_df = otomatik_hesaplari_uygula(yeni_df, secili_sablon)
    st.rerun()

# =========================================================
# TABLO EDİTÖRÜ (KULLANICI ARAYÜZÜ)
# =========================================================
df_ui = st.session_state.veri_df.copy()
# Kullanıcıya görsel olarak daha anlaşılır olması için index (sıra no) 1'den başlatılır
df_ui.index = df_ui.index + 1 

# Tablodaki sütunların veri tiplerine göre arayüzdeki input (girdi) şekillerini ayarlar
col_config = {}
for col in df_ui.columns:
    low = str(col).strip().lower()
    if low in {"adet", "qty", "quantity", "miktar"}:
        col_config[col] = st.column_config.NumberColumn(col, format="%d")
    elif low in {"birim fiyatı", "birim fiyati", "birim fiyat", "unit price", "price", "tutar", "total", "amount"}:
        col_config[col] = st.column_config.NumberColumn(col, format=f"%d {sembol}")
    else:
        col_config[col] = st.column_config.TextColumn(col)

st.info("💡 Tablodaki hücrelerin üzerine tıklayıp değiştirebilirsiniz. Yeni satır için '+' butonunu kullanın.")

# Streamlit data_editor bileşeni (Etkileşimli Excel benzeri yapı)
duzenlenmis_df_ui = st.data_editor(
    df_ui,
    column_config=col_config,
    num_rows="dynamic",
    use_container_width=True,
    hide_index=False, 
    key="veri_editoru"
)

# Arka planda işlem yapabilmek için indexi tekrar 0 tabanlı sisteme döndürürüz
duzenlenmis_df = duzenlenmis_df_ui.copy()
duzenlenmis_df.reset_index(drop=True, inplace=True)

# =========================================================
# OTOMATİK YENİLEME (TETİKLEYİCİ) SİSTEMİ
# Yeni bir satır eklendiğinde veya veri değiştiğinde hesaplamaları günceller
# =========================================================
hesaplanmis_df = otomatik_hesaplari_uygula(duzenlenmis_df, secili_sablon)

guncelle = False
if len(duzenlenmis_df) != len(st.session_state.veri_df):
    guncelle = True
else:
    for col in duzenlenmis_df.columns:
        low = str(col).lower().strip()
        if low == "kdv":
            s1 = st.session_state.veri_df[col].astype(str).fillna("").str.strip()
            s2 = hesaplanmis_df[col].astype(str).fillna("").str.strip()
            if not s1.equals(s2):
                guncelle = True
                break
        elif low in ["tutar", "total", "amount"]:
            s1 = pd.to_numeric(st.session_state.veri_df[col], errors='coerce').fillna(0).round(2)
            s2 = pd.to_numeric(hesaplanmis_df[col], errors='coerce').fillna(0).round(2)
            if not s1.equals(s2):
                guncelle = True
                break

if guncelle:
    st.session_state.veri_df = hesaplanmis_df
    st.rerun() # Sayfayı yeni hesaplanmış değerlerle yenile

duzenlenmis_df = hesaplanmis_df

# =========================================================
# TOPLAM HESAPLAMALARI (ARA TOPLAM, KDV, GENEL TOPLAM)
# =========================================================
toplam_sutunu = toplam_sutununu_bul(duzenlenmis_df, secili_sablon)
fiyatlar = pd.to_numeric(duzenlenmis_df[toplam_sutunu], errors='coerce').fillna(0)

ara_toplam = fiyatlar.sum()

# Her satırın KDV oranını bağımsız olarak işleme alır
if secili_sablon == "📄 Standart Proforma Fatura":
    kdv_col = kolonu_bul(duzenlenmis_df.columns, {"kdv"})
    if kdv_col is not None:
        def kdv_parse(val):
            try:
                if pd.isna(val): return 0.20
                v = str(val).lower().replace('%', '').replace(',', '.').strip()
                if not v or v == 'nan' or v == 'none': return 0.20
                return float(v) / 100.0
            except Exception:
                return 0.20
                
        satir_kdv_oranlari = duzenlenmis_df[kdv_col].apply(kdv_parse)
        kdv = (fiyatlar * satir_kdv_oranlari).sum()
    else:
        kdv = ara_toplam * 0.20
else:
    kdv = ara_toplam * 0.20

genel_toplam = ara_toplam + kdv

# Ekranda gösterilecek formatlı string metinleri oluşturur
ara_str = f"{ara_toplam:,.0f}".replace(",", ".") + f" {kur_metin}"
kdv_str = f"{kdv:,.0f}".replace(",", ".") + f" {kur_metin}"
genel_str = f"{genel_toplam:,.0f}".replace(",", ".") + f" {kur_metin}"

# Görüntüleme Paneli (Metrikler)
st.write("---")
col_a, col_b, col_c = st.columns(3)
col_a.metric("Ara Toplam", ara_str)
col_b.metric("KDV Tutarı", kdv_str)
col_c.metric("Genel Toplam", genel_str)
st.write("---")

st.subheader("📄 Belge Altı Notları")
st.text_area(
    "Bu alana yazdığınız metin belgenin altına eklenecektir:",
    key="not_alani",
    height=150,
    placeholder="Buraya notlarınızı veya banka hesap bilgilerinizi girebilirsiniz..."
)

if st.button("🔄 Notları Kaydet"):
    st.success("Notlarınız başarıyla hafızaya alındı! Çıktı alabilirsiniz.")

st.write("---")

# --- TASLAK KAYDETME (İNDİRME) ALANI ---
st.sidebar.markdown("---")
# Mevcut tüm verileri tek bir JSON objesinde toplayıp dışa aktarır
yedek_verisi = {
    "sablon": secili_sablon,
    "teklif_basligi": teklif_basligi_str,
    "notlar": st.session_state.not_alani,
    "df": duzenlenmis_df.to_dict(orient="records")
}
yedek_json = json.dumps(yedek_verisi, indent=4, ensure_ascii=False)

st.sidebar.download_button(
    label="💾 Mevcut Taslağı İndir (.json)",
    data=yedek_json,
    file_name=f"Innomarin_Taslak_{dosya_tarihi}.json",
    mime="application/json",
    use_container_width=True
)

# =========================================================
# ÇIKTI MOTORU 1: WORD DOKÜMANI OLUŞTURMA
# =========================================================
def word_olustur(dataframe, a_str, k_str, g_str, tarih, notlar, kur_m, sablon_tipi, gizle_aktif, teklif_basligi):
    """
    Girilen verileri işleyerek Microsoft Word (.docx) dosyası üretir.
    Logoyu Üst Bilgi (Header) kısmına yerleştirerek sayfa hizalamasını korur.
    """
    df_out = dataframe.copy()
    if gizle_aktif:
        birim_sutun = get_birim_col(df_out.columns)
        if birim_sutun:
            df_out = df_out.drop(columns=[birim_sutun])
            
    headers = ['Sıra' if sablon_tipi != "⚓ INNOMAR Özel Teklif" else 'ITEM NO'] + list(df_out.columns)

    if os.path.exists(WORD_TEMPLATE):
        doc = Document(WORD_TEMPLATE)
    else:
        doc = Document()

    # Sayfa Marjları ve Header Konfigürasyonu
    for section in doc.sections:
        section.top_margin = Cm(2.0)       
        section.bottom_margin = Cm(4.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.0)  
        
        # Logoyu Header'a (Üst Bilgiye) Ekleme İşlemi
        if os.path.exists("ust_bar.png"):
            header = section.header
            p_logo = header.paragraphs[0] if len(header.paragraphs) > 0 else header.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.text = "" 
            r_logo = p_logo.add_run()
            r_logo.add_picture("ust_bar.png", width=Cm(17))

    # Tablonun logoya girmemesi için alt satıra itme işlemi
    doc.add_paragraph("\n\n\n")

    if sablon_tipi == "⚓ INNOMAR Özel Teklif":
        p_title = doc.add_paragraph()
        p_title.add_run(f"•   {teklif_basligi}").bold = True
    else:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.add_run("PROFORMA FATURA").bold = True

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.add_run(f"TARİH / DATE: {tarih}").bold = True

    table = doc.add_table(rows=1, cols=len(headers))
    guvenli_table_style(table, "Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Tablo Sütun Başlıkları Çizimi
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)

        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        align = get_alignment(header)
        if align == 'R':
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'C':
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tablo İçeriklerinin Doldurulması
    for index, row in df_out.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(index + 1)
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for c_idx, col_name in enumerate(df_out.columns):
            val = row[col_name]
            align = get_alignment(col_name)
            low = str(col_name).strip().lower()

            # Veri Formatlama Mantığı
            if low in {"price", "tutar", "total", "amount", "birim fiyatı", "birim fiyati", "birim fiyat", "unit price"}:
                text_val = format_money_value(val, kur_m)
            elif low in {"adet", "qty", "quantity", "miktar", "unit"}:
                # Sayıları tam sayıya çevirme filtresi ("2.0" yerine "2")
                try:
                    f_val = float(val)
                    text_val = str(int(f_val)) if f_val.is_integer() else str(val)
                except:
                    text_val = str(val)
            else:
                text_val = str(val)

            row_cells[c_idx + 1].text = text_val

            if row_cells[c_idx + 1].paragraphs and row_cells[c_idx + 1].paragraphs[0].runs:
                row_cells[c_idx + 1].paragraphs[0].runs[0].font.size = Pt(9)

            if align == 'R':
                row_cells[c_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'C':
                row_cells[c_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tablo Altı Toplamlar Bölümü
    labels = (
        ["TOTAL PRICE", "VAT (20%)", "GRAND TOTAL"]
        if sablon_tipi == "⚓ INNOMAR Özel Teklif"
        else ["Ara Toplam", "KDV", "GENEL TOPLAM"]
    )
    values = [a_str, k_str, g_str]

    for i in range(3):
        row_cells = table.add_row().cells
        
        # Sola boş kalan hücreleri estetik amaçlı birleştir (Merge)
        if len(headers) > 2:
            row_cells[0].merge(row_cells[-3])
            row_cells[0].text = ""
            
        row_cells[-2].text = ""
        r1 = row_cells[-2].paragraphs[0].add_run(labels[i])
        r1.bold = True
        r1.font.size = Pt(9)
        row_cells[-2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        row_cells[-1].text = ""
        r2 = row_cells[-1].paragraphs[0].add_run(values[i])
        r2.bold = True
        r2.font.size = Pt(9)
        row_cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Notların Eklenmesi
    for satir in str(notlar).split('\n'):
        p = doc.add_paragraph(satir)
        if p.runs:
            p.runs[0].font.size = Pt(9)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =========================================================
# ÇIKTI MOTORU 2: PDF DOKÜMANI OLUŞTURMA
# =========================================================
def pdf_olustur(dataframe, a_str, k_str, g_str, tarih, notlar, kur_m, sablon_tipi, gizle_aktif, teklif_basligi):
    """
    Verilen verilerle otomatik sayfa atlama, sayfa sonu algılama ve başlık tekrarlama
    mekanizmalarına sahip kurumsal bir PDF dokümanı oluşturur.
    """
    df_out = dataframe.copy()
    if gizle_aktif:
        birim_sutun = get_birim_col(df_out.columns)
        if birim_sutun:
            df_out = df_out.drop(columns=[birim_sutun])

    class PDF(FPDF):
        def header(self):
            # Her yeni sayfa eklendiğinde arka plan logosunu otomatik basar
            if os.path.exists(ANTET_DOSYASI):
                self.image(ANTET_DOSYASI, x=0, y=0, w=210, h=297)
            # Metnin yeni sayfada logoyu ezmemesi için başlangıç noktasını 85mm'ye ayarlar
            self.set_y(85)

    pdf = PDF()
    pdf.add_page()
    # Alt marjini 45mm yaparak antetin alt kısmındaki tasarıma yazı binmesini engeller
    pdf.set_auto_page_break(auto=True, margin=45)

    pdf.set_font('Arial', 'B', 10)
    pdf.set_text_color(0, 0, 0)

    # Başlık ve Tarih Yerleşimi (Fatura Başlığı daha yukarıda başlatılır)
    if sablon_tipi == "⚓ INNOMAR Özel Teklif":
        pdf.set_y(65)
        pdf.cell(130, 8, chr(149) + f'   {teklif_basligi}', 0, 0, 'L')
        pdf.cell(60, 8, f'DATE: {tarih}', 0, 1, 'R')
    else:
        pdf.set_y(65)
        pdf.set_font('Arial', 'B', 13)
        pdf.cell(0, 8, 'PROFORMA FATURA', 0, 1, 'C')
        
        pdf.set_y(75)
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, f'TARIH: {tarih}', 0, 1, 'R')

    pdf.ln(4)

    headers = ['Sıra' if sablon_tipi != "⚓ INNOMAR Özel Teklif" else 'NO'] + list(df_out.columns)
    widths = get_pdf_widths(headers)

    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.2)
    
    def tablo_basliklarini_ciz():
        """Tablonun en üst kısmındaki ana başlıkları (Sıra, Açıklama, Adet vb.) çizer."""
        pdf.set_fill_color(255, 255, 255)
        pdf.set_font('Arial', 'B', 9)
        for idx, header in enumerate(headers):
            pdf.cell(widths[idx], 10, cevir_tr(str(header)), 1, 0, 'C', fill=True)
        pdf.ln()

    tablo_basliklarini_ciz()

    pdf.set_font('Arial', '', 8)
    
    # Satırların Çizilmesi ve Sayfa Sonu (Page Break) Kontrolü
    for index, row in df_out.iterrows():
        # EĞER Y EKSENİ (DİKEY POZİSYON) 240MM'Yİ GEÇTİYSE SAYFA DOLDU DEMEKTİR
        # Yeni sayfa açıp tablo başlıklarını tekrar yazdırarak kopukluğu engeller
        if pdf.get_y() > 240:
            pdf.add_page()
            tablo_basliklarini_ciz()
            pdf.set_font('Arial', '', 8)

        pdf.cell(widths[0], 8, str(index + 1), 1, 0, 'C', fill=True)

        for c_idx, col_name in enumerate(df_out.columns):
            val = row[col_name]
            align = get_alignment(col_name)
            low = str(col_name).strip().lower()

            if low in {"price", "tutar", "total", "amount", "birim fiyatı", "birim fiyati", "birim fiyat", "unit price"}:
                yazilacak = format_money_value(val, kur_m)
            elif low in {"adet", "qty", "quantity", "miktar", "unit"}:
                try:
                    f_val = float(val)
                    yazilacak = cevir_tr(str(int(f_val)) if f_val.is_integer() else str(val))
                except:
                    yazilacak = cevir_tr(str(val))
            else:
                yazilacak = cevir_tr(str(val))

            pdf.cell(widths[c_idx + 1], 8, yazilacak, 1, 0, align, fill=True)

        pdf.ln()

    # Toplamların Kesilmemesi İçin Kontrol (Eğer yer kalmadıysa toplamları yeni sayfaya taşır)
    if pdf.get_y() > 225:
        pdf.add_page()

    w_val = widths[-1]

    # Alt Toplamlar İçin Hizalama Hesabı
    label_w = 0
    label_cols = 0
    for w in reversed(widths[:-1]):
        label_w += w
        label_cols += 1
        if label_w >= 35:
            break

    w_label = label_w
    w_empty = sum(widths[:-1 - label_cols]) if len(widths) > label_cols + 1 else 0

    # Toplamlar Tablosu Çizimi
    pdf.set_font('Arial', '', 9)
    if w_empty > 0:
        pdf.cell(w_empty, 8, '', 0, 0)
    pdf.cell(w_label, 8, 'Ara Toplam' if sablon_tipi != "⚓ INNOMAR Özel Teklif" else 'TOTAL PRICE', 1, 0, 'R')
    pdf.set_font('Arial', 'B', 9)
    pdf.cell(w_val, 8, a_str, 1, 1, 'R')

    pdf.set_font('Arial', '', 9)
    if w_empty > 0:
        pdf.cell(w_empty, 8, '', 0, 0)
    pdf.cell(w_label, 8, 'KDV' if sablon_tipi != "⚓ INNOMAR Özel Teklif" else 'VAT (20%)', 1, 0, 'R')
    pdf.set_font('Arial', 'B', 9)
    pdf.cell(w_val, 8, k_str, 1, 1, 'R')

    pdf.set_font('Arial', 'B', 9)
    if w_empty > 0:
        pdf.cell(w_empty, 8, '', 0, 0)
    pdf.cell(w_label, 8, 'GENEL TOPLAM' if sablon_tipi != "⚓ INNOMAR Özel Teklif" else 'GRAND TOTAL', 1, 0, 'R')
    pdf.cell(w_val, 8, g_str, 1, 1, 'R')

    pdf.ln(10)
    pdf.set_font('Arial', '', 8)
    pdf.multi_cell(0, 5, cevir_tr(notlar))

    return pdf.output(dest='S').encode('latin-1')

# =========================================================
# ÇIKTI MOTORU 3: EXCEL DOKÜMANI OLUŞTURMA
# =========================================================
def excel_olustur(dataframe, a_str, k_str, g_str, tarih, notlar, kur_m, sablon_tipi, gizle_aktif, teklif_basligi):
    """
    Pandas verilerini saf Excel formatına (XLSX) çevirir.
    Hücre borderları (çizgiler), renk paletleri ve image injection işlemlerini yapar.
    """
    df_out = dataframe.copy()
    if gizle_aktif:
        birim_sutun = get_birim_col(df_out.columns)
        if birim_sutun:
            df_out = df_out.drop(columns=[birim_sutun])
            
    wb = Workbook()
    ws = wb.active
    ws.title = "Belge"
    ws.sheet_view.showGridLines = True

    row_idx = 1

    # Excel içine Logoyu Gömme İşlemi (Image Injection)
    if os.path.exists(ANTET_DOSYASI):
        img = xlImage(ANTET_DOSYASI)
        oran = 760 / img.width
        img.width = 760
        img.height = int(img.height * oran)
        ws.add_image(img, 'A1')
        # Logonun hücreleri ezmemesi için başlangıç satırını 22'ye öteliyoruz
        row_idx = 22
    else:
        ws['A1'] = "UYARI: antet.png bulunamadı."
        ws['A1'].font = Font(bold=True, color="FF0000")
        row_idx = 6

    # Başlık ve Müşteri Bilgileri Yapılandırması
    if sablon_tipi == "⚓ INNOMAR Özel Teklif":
        ws.sheet_view.showGridLines = False
        ws[f'B{row_idx}'] = f"• {teklif_basligi}"
        ws[f'B{row_idx}'].font = Font(bold=True, size=12)
        ws.cell(row=row_idx, column=len(df_out.columns) + 2).value = f"DATE: {tarih}"
        ws.cell(row=row_idx, column=len(df_out.columns) + 2).font = Font(bold=True)
        row_idx += 2
    else:
        ws[f'C{row_idx}'] = "PROFORMA FATURA"
        ws[f'C{row_idx}'].font = Font(bold=True, size=16)
        ws[f'C{row_idx}'].alignment = Alignment(horizontal="center")

        tarih_sutun = len(df_out.columns) - 1
        if tarih_sutun < 4:
            tarih_sutun = 4

        row_idx += 2
        ws.cell(row=row_idx, column=1).value = "Fatura Kesilen:"
        ws.cell(row=row_idx, column=1).font = Font(bold=True, size=12)
        ws.cell(row=row_idx, column=tarih_sutun).value = "Fatura Numarası: ........................"
        ws.cell(row=row_idx, column=tarih_sutun).font = Font(size=12)

        row_idx += 1
        ws.cell(row=row_idx, column=1).value = "Müşteri Adı: .............................................................."
        ws.cell(row=row_idx, column=1).font = Font(size=12)
        ws.cell(row=row_idx, column=tarih_sutun).value = "Fatura Tarihi:"
        ws.cell(row=row_idx, column=tarih_sutun).font = Font(bold=True, size=12)
        ws.cell(row=row_idx, column=tarih_sutun + 1).value = tarih
        ws.cell(row=row_idx, column=tarih_sutun + 1).font = Font(bold=True, size=12)

        row_idx += 1
        ws.cell(row=row_idx, column=1).value = "Adres: ........................................................................"
        ws.cell(row=row_idx, column=1).font = Font(size=12)

        row_idx += 2

    headers = ['Sıra' if sablon_tipi != "⚓ INNOMAR Özel Teklif" else 'ITEM NO'] + list(df_out.columns)
    set_excel_col_widths(ws, headers)

    thin_border = Border(
        left=Side(border_style='thin', color='000000'),
        right=Side(border_style='thin', color='000000'),
        top=Side(border_style='thin', color='000000'),
        bottom=Side(border_style='thin', color='000000')
    )

    bg_fill = (
        PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        if sablon_tipi == "⚓ INNOMAR Özel Teklif"
        else PatternFill(start_color="EBE4D5", end_color="EBE4D5", fill_type="solid")
    )

    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=row_idx, column=col_num)
        cell.value = str(header)
        cell.font = Font(bold=True)
        cell.border = thin_border
        cell.fill = bg_fill

        align = get_alignment(header)
        if align == 'R':
            cell.alignment = Alignment(horizontal="right")
        elif align == 'C':
            cell.alignment = Alignment(horizontal="center")
        else:
            cell.alignment = Alignment(horizontal="left")
    row_idx += 1

    for index, row in df_out.iterrows():
        cell = ws.cell(row=row_idx, column=1)
        cell.value = index + 1
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

        for c_idx, col_name in enumerate(df_out.columns):
            val = row[col_name]
            cell = ws.cell(row=row_idx, column=c_idx + 2)
            align = get_alignment(col_name)
            low = str(col_name).strip().lower()

            if low in {"price", "tutar", "total", "amount", "birim fiyatı", "birim fiyati", "birim fiyat", "unit price"}:
                cell.value = format_money_value(val, kur_m)
                if align == 'R':
                    cell.alignment = Alignment(horizontal="right")
                elif align == 'C':
                    cell.alignment = Alignment(horizontal="center")
                else:
                    cell.alignment = Alignment(horizontal="left")
            elif low in {"adet", "qty", "quantity", "miktar", "unit"}:
                try:
                    f_val = float(val)
                    cell.value = str(int(f_val)) if f_val.is_integer() else str(val)
                except:
                    cell.value = str(val)
                if align == 'R':
                    cell.alignment = Alignment(horizontal="right")
                elif align == 'C':
                    cell.alignment = Alignment(horizontal="center")
                else:
                    cell.alignment = Alignment(horizontal="left")
            else:
                cell.value = str(val)
                if align == 'R':
                    cell.alignment = Alignment(horizontal="right")
                elif align == 'C':
                    cell.alignment = Alignment(horizontal="center")
                else:
                    cell.alignment = Alignment(horizontal="left")

            cell.border = thin_border

        row_idx += 1

    tot_col = len(headers) - 1
    val_col = len(headers)

    ws.cell(row=row_idx, column=tot_col).value = "Ara Toplam" if sablon_tipi != "⚓ INNOMAR Özel Teklif" else "TOTAL PRICE"
    ws.cell(row=row_idx, column=tot_col).border = thin_border
    ws.cell(row=row_idx, column=val_col).value = a_str
    ws.cell(row=row_idx, column=val_col).font = Font(bold=True)
    ws.cell(row=row_idx, column=val_col).border = thin_border
    ws.cell(row=row_idx, column=val_col).alignment = Alignment(horizontal="right")
    row_idx += 1

    ws.cell(row=row_idx, column=tot_col).value = "KDV" if sablon_tipi != "⚓ INNOMAR Özel Teklif" else "VAT (20%)"
    ws.cell(row=row_idx, column=tot_col).border = thin_border
    ws.cell(row=row_idx, column=val_col).value = k_str
    ws.cell(row=row_idx, column=val_col).font = Font(bold=True)
    ws.cell(row=row_idx, column=val_col).border = thin_border
    ws.cell(row=row_idx, column=val_col).alignment = Alignment(horizontal="right")
    row_idx += 1

    ws.cell(row=row_idx, column=tot_col).value = "GENEL TOPLAM" if sablon_tipi != "⚓ INNOMAR Özel Teklif" else "GRAND TOTAL"
    ws.cell(row=row_idx, column=tot_col).border = thin_border
    ws.cell(row=row_idx, column=val_col).value = g_str
    ws.cell(row=row_idx, column=val_col).font = Font(bold=True)
    ws.cell(row=row_idx, column=val_col).border = thin_border
    ws.cell(row=row_idx, column=val_col).alignment = Alignment(horizontal="right")
    row_idx += 2

    for satir in str(notlar).split('\n'):
        ws[f'A{row_idx}'] = satir
        row_idx += 1

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()

# =========================================================
# DOSYA DURUMU BİLGİLENDİRME PANELİ
# =========================================================
st.markdown("### 📎 Şablon Durumu")

if os.path.exists(WORD_TEMPLATE):
    st.success("✅ word_template.docx bulundu. Word çıktısı template üzerinden oluşturulacak.")
else:
    st.warning("⚠️ word_template.docx bulunamadı. Word çıktısı düz oluşur.")

if os.path.exists(ANTET_DOSYASI):
    st.success("✅ antet.png bulundu. PDF ve Excel çıktılarında kullanılacak.")
else:
    st.warning("⚠️ antet.png bulunamadı. PDF ve Excel çıktılarında antet görünmeyecek.")

# =========================================================
# İNDİRME BUTONLARI (BİTİŞ BÖLÜMÜ)
# =========================================================
st.markdown("### 📥 Çıktı Al")
btn_word, btn_excel, btn_pdf = st.columns(3)

with btn_word:
    st.download_button(
        "📄 WORD İNDİR",
        data=word_olustur(
            duzenlenmis_df,
            ara_str,
            kdv_str,
            genel_str,
            tarih_metni,
            st.session_state.not_alani,
            kur_metin,
            secili_sablon,
            gizle_checkbox,
            teklif_basligi_str
        ),
        file_name=f"{secili_sablon.split()[1]}_{dosya_tarihi}.docx",
        type="primary",
        use_container_width=True
    )

with btn_excel:
    st.download_button(
        "📊 EXCEL İNDİR",
        data=excel_olustur(
            duzenlenmis_df,
            ara_str,
            kdv_str,
            genel_str,
            tarih_metni,
            st.session_state.not_alani,
            kur_metin,
            secili_sablon,
            gizle_checkbox,
            teklif_basligi_str
        ),
        file_name=f"{secili_sablon.split()[1]}_{dosya_tarihi}.xlsx",
        type="primary",
        use_container_width=True
    )

with btn_pdf:
    st.download_button(
        "📕 PDF İNDİR",
        data=pdf_olustur(
            duzenlenmis_df,
            ara_str,
            kdv_str,
            genel_str,
            tarih_metni,
            st.session_state.not_alani,
            kur_metin,
            secili_sablon,
            gizle_checkbox,
            teklif_basligi_str
        ),
        file_name=f"{secili_sablon.split()[1]}_{dosya_tarihi}.pdf",
        type="primary",
        use_container_width=True
    )
